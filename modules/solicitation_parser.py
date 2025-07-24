"""
Solicitation parsing module for extracting structured data from PDF, Word, and text documents.
Follows existing module patterns with stateless utility class design.
Includes advanced features: caching, templates, logging, and multi-format support.
"""

import re
import time
import json
import logging
import psutil
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from datetime import datetime
import PyPDF2
import streamlit as st

# Optional dependencies with graceful fallback
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .data_models import Solicitation, ExtractionConfig, SolicitationParsingResult


class SolicitationParser:
    """
    Handles multi-format solicitation parsing with advanced features:
    - PDF, Word, and text document support
    - Template system for extraction configurations
    - Comprehensive logging and monitoring
    - Document type auto-detection
    - Performance optimization with caching
    """
    
    def __init__(self, templates_dir: str = "./data/templates", logs_dir: str = "./data/logs"):
        """Initialize parser with template and logging support."""
        self.templates_dir = Path(templates_dir)
        self.logs_dir = Path(logs_dir)
        
        # Ensure directories exist
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        self.logs_dir.mkdir(parents=True, exist_ok=True)
        
        # Setup logging
        self.logger = self._setup_logging()
        
        # Load extraction configurations
        self.extraction_configs = self._load_default_configs()
        
        # Document type patterns for auto-detection
        self.document_type_patterns = self._load_document_type_patterns()
        
        # Performance monitoring
        self.performance_stats = {
            'total_documents_processed': 0,
            'processing_times': [],
            'memory_usage': [],
            'error_count': 0,
            'success_count': 0
        }
    
    def _setup_logging(self) -> logging.Logger:
        """Setup comprehensive logging for parsing operations."""
        logger = logging.getLogger('solicitation_parser')
        logger.setLevel(logging.INFO)
        
        # Avoid duplicate handlers
        if logger.handlers:
            return logger
        
        # File handler for detailed logs
        log_file = self.logs_dir / f"parsing_{datetime.now().strftime('%Y%m%d')}.log"
        file_handler = logging.FileHandler(log_file)
        file_handler.setLevel(logging.INFO)
        
        # Console handler for errors
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.ERROR)
        
        # Formatter
        formatter = logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        
        logger.addHandler(file_handler)
        logger.addHandler(console_handler)
        
        return logger
    
    def _load_document_type_patterns(self) -> Dict[str, Dict[str, Any]]:
        """Load document type detection patterns."""
        return {
            'nsf': {
                'patterns': [
                    r'NSF\s+\d{2}-\d{3}',
                    r'National Science Foundation',
                    r'Program Solicitation',
                    r'Cognizant Program Officer'
                ],
                'confidence_threshold': 2,
                'extraction_config': 'nsf_template'
            },
            'nih': {
                'patterns': [
                    r'NIH\s+[A-Z]{2,4}-\d+',
                    r'National Institutes of Health',
                    r'Funding Opportunity Announcement',
                    r'FOA Number'
                ],
                'confidence_threshold': 2,
                'extraction_config': 'nih_template'
            },
            'generic_rfp': {
                'patterns': [
                    r'Request for Proposals?',
                    r'RFP\s+\d+',
                    r'Proposal Submission',
                    r'Funding Opportunity'
                ],
                'confidence_threshold': 1,
                'extraction_config': 'generic_template'
            }
        }
    
    @st.cache_data
    def extract_text_from_file(_self, file_path: str) -> tuple[str, str]:
        """
        Extract text from various file formats with caching.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, detected_file_type)
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        _self.logger.info(f"Extracting text from {file_path} (type: {file_extension})")
        
        try:
            if file_extension == '.pdf':
                text = _self._extract_text_from_pdf(str(file_path))
                return text, 'pdf'
            elif file_extension == '.docx':
                text = _self._extract_text_from_docx(str(file_path))
                return text, 'docx'
            elif file_extension in ['.txt', '.text']:
                text = _self._extract_text_from_txt(str(file_path))
                return text, 'txt'
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            _self.logger.error(f"Failed to extract text from {file_path}: {str(e)}")
            raise
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using PyPDF2."""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
    
    def _extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word document support")
        
        doc = DocxDocument(docx_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    def _extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text from plain text file."""
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read().strip()
    
    def detect_document_type(self, text: str) -> tuple[str, float]:
        """
        Auto-detect document type based on content patterns.
        
        Args:
            text: Document text content
            
        Returns:
            Tuple of (document_type, confidence_score)
        """
        type_scores = {}
        
        for doc_type, config in self.document_type_patterns.items():
            score = 0
            for pattern in config['patterns']:
                matches = len(re.findall(pattern, text, re.IGNORECASE))
                score += matches
            
            type_scores[doc_type] = score
        
        # Find best match
        best_type = max(type_scores, key=type_scores.get)
        best_score = type_scores[best_type]
        
        # Check if confidence threshold is met
        threshold = self.document_type_patterns[best_type]['confidence_threshold']
        confidence = min(1.0, best_score / threshold) if threshold > 0 else 0.0
        
        self.logger.info(f"Document type detection: {best_type} (confidence: {confidence:.2f})")
        
        return best_type, confidence   
 
    def parse_document(self, file_path: str, template_name: Optional[str] = None) -> SolicitationParsingResult:
        """
        Main parsing method for multi-format documents with advanced features.
        
        Args:
            file_path: Path to the document file
            template_name: Optional template name to use for extraction
            
        Returns:
            SolicitationParsingResult with extracted data and metadata
        """
        start_time = time.time()
        memory_before = psutil.Process().memory_info().rss / 1024 / 1024  # MB
        
        try:
            # Validate input
            if not file_path or not Path(file_path).exists():
                raise Exception(f"File not found: {file_path}")
            
            # Extract text with caching
            text, file_type = self.extract_text_from_file(file_path)
            
            if not text or len(text.strip()) < 50:
                raise Exception("Document appears to be empty or contains insufficient text")
            
            # Auto-detect document type if no template specified
            if not template_name:
                doc_type, type_confidence = self.detect_document_type(text)
                if type_confidence > 0.5:
                    template_name = self.document_type_patterns[doc_type]['extraction_config']
                    self.logger.info(f"Auto-selected template: {template_name}")
            
            # Load template if specified
            if template_name:
                self._load_template(template_name)
            
            # Extract structured data
            extracted_data = self.extract_solicitation_data(text)
            
            # Validate and clean
            extracted_data = self._validate_and_clean_data(extracted_data)
            
            # Calculate quality metrics
            confidence_score, missing_fields = self._assess_quality(extracted_data)
            
            # Performance monitoring
            processing_time = time.time() - start_time
            memory_after = psutil.Process().memory_info().rss / 1024 / 1024  # MB
            memory_used = memory_after - memory_before
            
            self._update_performance_stats(processing_time, memory_used, success=True)
            
            # Log successful processing
            self.logger.info(
                f"Successfully processed {file_path} "
                f"(type: {file_type}, time: {processing_time:.2f}s, "
                f"memory: {memory_used:.1f}MB, confidence: {confidence_score:.2f})"
            )
            
            return SolicitationParsingResult(
                extracted_data=extracted_data,
                confidence_score=confidence_score,
                missing_fields=missing_fields,
                source_file=file_path,
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = time.time() - start_time
            self._update_performance_stats(processing_time, 0, success=False)
            
            self.logger.error(f"Failed to process {file_path}: {str(e)}")
            
            # Return partial results if possible
            try:
                if 'text' in locals() and text:
                    extracted_data = self.extract_solicitation_data(text)
                    extracted_data = self._validate_and_clean_data(extracted_data)
                    confidence_score, missing_fields = self._assess_quality(extracted_data)
                else:
                    extracted_data = {}
                    confidence_score = 0.0
                    missing_fields = list(self.extraction_configs.keys())
            except:
                extracted_data = {}
                confidence_score = 0.0
                missing_fields = list(self.extraction_configs.keys())
            
            return SolicitationParsingResult(
                extracted_data=extracted_data,
                confidence_score=confidence_score,
                missing_fields=missing_fields,
                source_file=file_path,
                processing_time=processing_time
            )
    
    def save_template(self, template_name: str, description: str = "") -> bool:
        """
        Save current extraction configuration as a reusable template.
        
        Args:
            template_name: Name for the template
            description: Optional description
            
        Returns:
            True if saved successfully
        """
        try:
            template_data = {
                'name': template_name,
                'description': description,
                'created_at': datetime.now().isoformat(),
                'extraction_configs': {
                    name: {
                        'field_name': config.field_name,
                        'patterns': config.patterns,
                        'extraction_type': config.extraction_type,
                        'post_process': config.post_process,
                        'required': config.required,
                        'default_value': config.default_value
                    }
                    for name, config in self.extraction_configs.items()
                }
            }
            
            template_file = self.templates_dir / f"{template_name}.json"
            with open(template_file, 'w') as f:
                json.dump(template_data, f, indent=2)
            
            self.logger.info(f"Saved template: {template_name}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to save template {template_name}: {str(e)}")
            return False
    
    def load_template(self, template_name: str) -> bool:
        """
        Load extraction configuration from a saved template.
        
        Args:
            template_name: Name of the template to load
            
        Returns:
            True if loaded successfully
        """
        return self._load_template(template_name)
    
    def _load_template(self, template_name: str) -> bool:
        """Internal method to load template."""
        try:
            template_file = self.templates_dir / f"{template_name}.json"
            
            if not template_file.exists():
                self.logger.warning(f"Template not found: {template_name}")
                return False
            
            with open(template_file, 'r') as f:
                template_data = json.load(f)
            
            # Convert back to ExtractionConfig objects
            self.extraction_configs = {}
            for name, config_data in template_data['extraction_configs'].items():
                self.extraction_configs[name] = ExtractionConfig(
                    field_name=config_data['field_name'],
                    patterns=config_data['patterns'],
                    extraction_type=config_data.get('extraction_type', 'single'),
                    post_process=config_data.get('post_process'),
                    required=config_data.get('required', False),
                    default_value=config_data.get('default_value', "")
                )
            
            self.logger.info(f"Loaded template: {template_name}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to load template {template_name}: {str(e)}")
            return False
    
    def list_templates(self) -> List[Dict[str, Any]]:
        """
        List all available templates with metadata.
        
        Returns:
            List of template information dictionaries
        """
        templates = []
        
        for template_file in self.templates_dir.glob("*.json"):
            try:
                with open(template_file, 'r') as f:
                    template_data = json.load(f)
                
                templates.append({
                    'name': template_data.get('name', template_file.stem),
                    'description': template_data.get('description', ''),
                    'created_at': template_data.get('created_at', ''),
                    'field_count': len(template_data.get('extraction_configs', {}))
                })
                
            except Exception as e:
                self.logger.error(f"Error reading template {template_file}: {str(e)}")
                continue
        
        return sorted(templates, key=lambda x: x['created_at'], reverse=True)
    
    def get_performance_stats(self) -> Dict[str, Any]:
        """Get comprehensive performance statistics."""
        stats = self.performance_stats.copy()
        
        if stats['processing_times']:
            stats['avg_processing_time'] = sum(stats['processing_times']) / len(stats['processing_times'])
            stats['max_processing_time'] = max(stats['processing_times'])
            stats['min_processing_time'] = min(stats['processing_times'])
        
        if stats['memory_usage']:
            stats['avg_memory_usage'] = sum(stats['memory_usage']) / len(stats['memory_usage'])
            stats['max_memory_usage'] = max(stats['memory_usage'])
        
        stats['success_rate'] = (
            stats['success_count'] / (stats['success_count'] + stats['error_count'])
            if (stats['success_count'] + stats['error_count']) > 0 else 0
        )
        
        return stats
    
    def _update_performance_stats(self, processing_time: float, memory_used: float, success: bool):
        """Update performance monitoring statistics."""
        self.performance_stats['total_documents_processed'] += 1
        self.performance_stats['processing_times'].append(processing_time)
        self.performance_stats['memory_usage'].append(memory_used)
        
        if success:
            self.performance_stats['success_count'] += 1
        else:
            self.performance_stats['error_count'] += 1
        
        # Keep only last 100 entries to prevent memory bloat
        if len(self.performance_stats['processing_times']) > 100:
            self.performance_stats['processing_times'] = self.performance_stats['processing_times'][-100:]
            self.performance_stats['memory_usage'] = self.performance_stats['memory_usage'][-100:]  
  
    def extract_solicitation_data(self, text: str) -> Dict[str, Any]:
        """
        Apply extraction patterns to get structured data.
        
        Args:
            text: Raw text content from document
            
        Returns:
            Dictionary with extracted field values
        """
        extracted_data = {}
        
        for field_name, config in self.extraction_configs.items():
            extracted_value = self._extract_field(text, config)
            extracted_data[field_name] = extracted_value
        
        return extracted_data
    
    def convert_to_solicitation(self, parsing_result: SolicitationParsingResult) -> Solicitation:
        """
        Convert extracted data to existing Solicitation model.
        
        Args:
            parsing_result: Result from document parsing process
            
        Returns:
            Solicitation object with extracted data
        """
        data = parsing_result.extracted_data
        
        # Normalize required skills to list format
        skills = data.get('required_skills', '')
        if isinstance(skills, str):
            # Split on common delimiters and clean up
            skills_list = [
                skill.strip() 
                for skill in re.split(r'[,;•\n\r]+', skills) 
                if skill.strip()
            ]
        else:
            skills_list = skills if isinstance(skills, list) else []
        
        return Solicitation(
            title=data.get('title', ''),
            abstract=data.get('abstract', ''),
            required_skills_checklist=skills_list,
            eligibility=data.get('eligibility'),
            description=data.get('description'),
            funding_amount=data.get('funding_amount'),
            deadline=data.get('deadline'),
            program=data.get('program'),
            parsing_metadata={
                'source_file': parsing_result.source_file,
                'processing_time': parsing_result.processing_time,
                'missing_fields': parsing_result.missing_fields
            },
            extraction_confidence=parsing_result.confidence_score
        )
    
    def _load_default_configs(self) -> Dict[str, ExtractionConfig]:
        """
        Load default extraction configurations for NSF solicitation fields.
        
        Returns:
            Dictionary mapping field names to extraction configurations
        """
        return {
            'title': ExtractionConfig(
                field_name='title',
                patterns=[
                    # NSF-specific title patterns
                    r'NSF\s+\d{2}-\d{3}:\s*([^\n]+(?:\n[^\n]+)*?)(?=\n(?:Broadening|Program Solicitation))',
                    r'Program Title:\s*([^\n]+)',
                    r'Solicitation Title:\s*([^\n]+)',
                    # Generic title patterns
                    r'Title:\s*([^\n]+)',
                ],
                required=True
            ),
            'abstract': ExtractionConfig(
                field_name='abstract',
                patterns=[
                    # Look for subtitle/description after main title
                    r'(?:NSF\s+\d{2}-\d{3}:[^\n]+\n[^\n]+\n)([^\n]+)(?=\nProgram Solicitation)',
                    r'Synopsis of Program:\s*(.*?)(?=\nCognizant Program Officer)',
                    r'Program Description\s*(.*?)(?=\n(?:III\.|Award Information))',
                    r'II\.\s*Program Description\s*(.*?)(?=\n(?:III\.|Award Information))',
                    # Generic patterns
                    r'Abstract:\s*(.*?)(?=\n\n|\n[A-Z][a-z]+:)',
                    r'Summary:\s*(.*?)(?=\n\n|\n[A-Z][a-z]+:)',
                    r'Overview:\s*(.*?)(?=\n\n|\n[A-Z][a-z]+:)',
                ],
                extraction_type='single',
                required=True
            ),
            'required_skills': ExtractionConfig(
                field_name='required_skills',
                patterns=[
                    # Look for collaboration and project types mentioned
                    r'enhancement of existing projects by virtue of new collaboration[^;]*;([^;]+(?:;[^;]+)*)',
                    r'initiation of new projects made possible by the collaboration[^;]*;([^;]+(?:;[^;]+)*)',
                    r'community building activities[^;]*;([^;]+(?:;[^;]+)*)',
                    # Generic skill patterns
                    r'Required.*?Skills?.*?:\s*(.*?)(?=\n\n|\n[A-Z][a-z]+:)',
                    r'Expertise.*?:\s*(.*?)(?=\n\n|\n[A-Z][a-z]+:)',
                    r'Qualifications.*?:\s*(.*?)(?=\n\n|\n[A-Z][a-z]+:)',
                    r'Technical.*?Requirements.*?:\s*(.*?)(?=\n\n|\n[A-Z][a-z]+:)'
                ],
                extraction_type='multiple',
                required=True
            ),
            'funding_amount': ExtractionConfig(
                field_name='funding_amount',
                patterns=[
                    # NSF-specific funding patterns
                    r'Anticipated Funding Amount:\s*\$?\s*([\d,]+(?:\.\d{2})?)',
                    r'Total Program Funding:\s*\$?\s*([\d,]+(?:\.\d{2})?)',
                    # Generic patterns
                    r'Budget:\s*\$?\s*([\d,]+(?:\.\d{2})?)',
                    r'Funding.*?Amount:\s*\$?\s*([\d,]+(?:\.\d{2})?)',
                    r'Award.*?Amount:\s*\$?\s*([\d,]+(?:\.\d{2})?)',
                    r'\$\s*([\d,]+(?:\.\d{2})?)'
                ],
                required=False
            ),
            'deadline': ExtractionConfig(
                field_name='deadline',
                patterns=[
                    # NSF submission window patterns
                    r'Submission Window Date\(s\)[^:]*:\s*([^\n]+(?:\n\s*[A-Z][a-z]+ \d+, \d+ - [A-Z][a-z]+ \d+, \d+)*)',
                    r'Due Dates?[^:]*:\s*([^\n]+)',
                    # Generic patterns
                    r'Deadline:\s*([^\n]+)',
                    r'Application.*?Deadline:\s*([^\n]+)',
                    r'Submission.*?Date:\s*([^\n]+)'
                ],
                required=False
            ),
            'program': ExtractionConfig(
                field_name='program',
                patterns=[
                    # NSF program number
                    r'(NSF\s+\d{2}-\d{3})',
                    r'Program Title:\s*([^\n]+)',
                    r'Program:\s*([^\n]+)',
                    r'Funding.*?Program:\s*([^\n]+)',
                    # NIH patterns for other agencies
                    r'NIH\s+([A-Z]{2,4}-\d+)'
                ],
                required=False
            ),
            'description': ExtractionConfig(
                field_name='description',
                patterns=[
                    # NSF program description section
                    r'II\.\s*Program Description\s*(.*?)(?=\nIII\.)',
                    r'Program Description\s*(.*?)(?=\n(?:III\.|Award Information))',
                    # Generic patterns
                    r'Background:\s*(.*?)(?=\n\n|\n[A-Z][a-z]+:)',
                    r'Objectives:\s*(.*?)(?=\n\n|\n[A-Z][a-z]+:)'
                ],
                extraction_type='single',
                required=False
            )
        }
    
    def _extract_field(self, text: str, config: ExtractionConfig) -> Any:
        """
        Extract a single field using the provided configuration.
        
        Args:
            text: Text to search in
            config: Extraction configuration
            
        Returns:
            Extracted value or default value if not found
        """
        for pattern in config.patterns:
            try:
                if config.extraction_type == 'multiple':
                    matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
                    if matches:
                        # Join multiple matches or return list
                        return ' '.join(matches) if isinstance(matches[0], str) else matches
                else:
                    match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                    if match:
                        return match.group(1).strip()
            except Exception:
                continue  # Try next pattern if current one fails
        
        return config.default_value
    
    def _validate_and_clean_data(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Validate and clean extracted data.
        
        Args:
            extracted_data: Raw extracted data
            
        Returns:
            Cleaned and validated data
        """
        cleaned_data = {}
        
        for field_name, value in extracted_data.items():
            if field_name not in self.extraction_configs:
                continue
                
            config = self.extraction_configs[field_name]
            
            # Clean and validate the value
            if value is None:
                cleaned_value = config.default_value
            elif isinstance(value, str):
                # Clean whitespace and normalize
                cleaned_value = ' '.join(value.split())
                
                # Remove common PDF artifacts
                cleaned_value = re.sub(r'\s+', ' ', cleaned_value)
                cleaned_value = re.sub(r'[^\w\s\-.,;:()\[\]$%]', ' ', cleaned_value)
                cleaned_value = cleaned_value.strip()
                
                # Apply field-specific validation
                if field_name == 'funding_amount':
                    # Ensure funding amount is numeric
                    cleaned_value = re.sub(r'[^\d,.]', '', cleaned_value)
                elif field_name == 'required_skills':
                    # Ensure skills are properly formatted
                    if cleaned_value:
                        # Split and clean individual skills
                        skills = [s.strip() for s in re.split(r'[;,\n]', cleaned_value) if s.strip()]
                        cleaned_value = '; '.join(skills[:10])  # Limit to 10 skills
                
                # Set to default if empty after cleaning
                if not cleaned_value:
                    cleaned_value = config.default_value
            else:
                cleaned_value = value
            
            cleaned_data[field_name] = cleaned_value
        
        return cleaned_data
    
    def _assess_quality(self, extracted_data: Dict[str, Any]) -> tuple[float, List[str]]:
        """
        Assess the quality of extracted data and identify missing fields.
        
        Args:
            extracted_data: Dictionary of extracted field values
            
        Returns:
            Tuple of (confidence_score, missing_fields)
        """
        total_fields = len(self.extraction_configs)
        required_fields = [
            name for name, config in self.extraction_configs.items() 
            if config.required
        ]
        
        # Check for missing required fields
        missing_fields = []
        filled_fields = 0
        quality_scores = []
        
        for field_name, config in self.extraction_configs.items():
            value = extracted_data.get(field_name, config.default_value)
            
            if value and str(value).strip():
                filled_fields += 1
                
                # Calculate field-specific quality score
                field_quality = self._calculate_field_quality(field_name, value)
                quality_scores.append(field_quality)
            elif config.required:
                missing_fields.append(field_name)
                quality_scores.append(0.0)
            else:
                quality_scores.append(0.5)  # Neutral score for optional empty fields
        
        # Calculate overall confidence score
        if quality_scores:
            avg_quality = sum(quality_scores) / len(quality_scores)
        else:
            avg_quality = 0.0
        
        # Penalize missing required fields more heavily
        required_penalty = len(missing_fields) * 0.3
        confidence_score = max(0, avg_quality - required_penalty)
        
        return confidence_score, missing_fields
    
    def _calculate_field_quality(self, field_name: str, value: Any) -> float:
        """
        Calculate quality score for a specific field.
        
        Args:
            field_name: Name of the field
            value: Extracted value
            
        Returns:
            Quality score between 0.0 and 1.0
        """
        if not value or not str(value).strip():
            return 0.0
        
        value_str = str(value).strip()
        
        # Field-specific quality assessment
        if field_name == 'title':
            # Good titles are descriptive and not too short/long
            if 10 <= len(value_str) <= 200:
                return 1.0
            elif len(value_str) < 10:
                return 0.3
            else:
                return 0.7
        
        elif field_name == 'abstract':
            # Good abstracts are substantial
            if len(value_str) >= 100:
                return 1.0
            elif len(value_str) >= 50:
                return 0.7
            else:
                return 0.4
        
        elif field_name == 'required_skills':
            # Good skills lists have multiple items
            skill_count = len([s for s in re.split(r'[;,\n]', value_str) if s.strip()])
            if skill_count >= 3:
                return 1.0
            elif skill_count >= 1:
                return 0.6
            else:
                return 0.2
        
        elif field_name == 'funding_amount':
            # Check if it looks like a valid amount
            if re.match(r'^\d{1,3}(,\d{3})*(\.\d{2})?$', value_str):
                return 1.0
            elif re.search(r'\d', value_str):
                return 0.6
            else:
                return 0.2
        
        elif field_name == 'program':
            # NSF program codes are high quality
            if re.match(r'NSF\s+\d{2}-\d{3}', value_str):
                return 1.0
            elif len(value_str) >= 5:
                return 0.7
            else:
                return 0.4
        
        else:
            # Generic quality assessment
            if len(value_str) >= 20:
                return 0.8
            elif len(value_str) >= 5:
                return 0.6
            else:
                return 0.3
    
    # Legacy methods for backward compatibility
    def parse_pdf_solicitation(self, pdf_path: str) -> SolicitationParsingResult:
        """Legacy method - use parse_document instead."""
        return self.parse_document(pdf_path)
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Legacy method - use extract_text_from_file instead."""
        text, _ = self.extract_text_from_file(pdf_path)
        return text